import os
import re
import io
import uuid
from typing import List, Dict, Tuple

from flask import Flask, render_template, request, jsonify
import pdfplumber
import docx
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 32 * 1024 * 1024  # 32 MB total per request

ALLOWED_EXTENSIONS = {"pdf", "docx"}

# A practical, recruiter-focused skill dictionary used for keyword matching.
# It is intentionally broad and lowercase. Multi-word skills are supported.
SKILL_DICTIONARY = [
    # Languages
    "python", "java", "c++", "c#", "javascript", "typescript", "go", "golang",
    "rust", "kotlin", "swift", "ruby", "php", "scala", "r", "matlab", "perl",
    "bash", "shell", "sql", "nosql", "html", "css", "sass", "less",
    # Web / frontend
    "react", "react.js", "next.js", "vue", "vue.js", "angular", "svelte",
    "redux", "tailwind", "tailwindcss", "bootstrap", "jquery", "webpack", "vite",
    # Backend / frameworks
    "node.js", "nodejs", "express", "express.js", "nestjs", "django", "flask",
    "fastapi", "spring", "spring boot", "rails", "ruby on rails", "laravel",
    "asp.net", ".net", "graphql", "rest", "rest api", "grpc",
    # Mobile
    "android", "ios", "flutter", "react native", "xamarin",
    # Data / ML
    "machine learning", "deep learning", "nlp", "natural language processing",
    "computer vision", "data science", "data analysis", "data engineering",
    "pandas", "numpy", "scipy", "scikit-learn", "sklearn", "tensorflow",
    "keras", "pytorch", "xgboost", "lightgbm", "spark", "pyspark", "hadoop",
    "hive", "kafka", "airflow", "dbt", "snowflake", "databricks", "tableau",
    "power bi", "looker", "matplotlib", "seaborn", "plotly",
    # Databases
    "postgresql", "postgres", "mysql", "mariadb", "sqlite", "mongodb", "redis",
    "cassandra", "dynamodb", "elasticsearch", "oracle", "sql server",
    # Cloud / DevOps
    "aws", "amazon web services", "azure", "gcp", "google cloud", "heroku",
    "docker", "kubernetes", "k8s", "terraform", "ansible", "jenkins", "gitlab ci",
    "github actions", "circleci", "ci/cd", "linux", "unix", "nginx", "apache",
    # Tools / methods
    "git", "github", "gitlab", "bitbucket", "jira", "confluence", "agile",
    "scrum", "kanban", "tdd", "bdd", "microservices", "serverless",
    # Soft / general
    "communication", "leadership", "problem solving", "teamwork",
    "project management", "stakeholder management", "mentoring",
    # Security
    "cybersecurity", "penetration testing", "owasp", "cryptography",
    # Misc
    "etl", "elt", "api", "oop", "design patterns", "system design",
    "data structures", "algorithms",
]


def allowed_file(filename: str) -> bool:
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_stream) -> str:
    text_chunks: List[str] = []
    with pdfplumber.open(file_stream) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
    return "\n".join(text_chunks)


def extract_text_from_docx(file_stream) -> str:
    document = docx.Document(file_stream)
    paragraphs = [p.text for p in document.paragraphs if p.text]
    # Also pull text from tables (resumes often use tables for layout)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def extract_text(filename: str, file_stream) -> str:
    ext = filename.rsplit(".", 1)[1].lower()
    if ext == "pdf":
        return extract_text_from_pdf(file_stream)
    if ext == "docx":
        return extract_text_from_docx(file_stream)
    return ""


def normalize_text(text: str) -> str:
    text = text.lower()
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def find_skills(text: str, skills: List[str]) -> List[str]:
    found = []
    lowered = " " + normalize_text(text) + " "
    for skill in skills:
        # word boundary match for short tokens; substring for multi-word
        if " " in skill or any(c in skill for c in ["+", "#", "."]):
            if skill in lowered:
                found.append(skill)
        else:
            pattern = r"\b" + re.escape(skill) + r"\b"
            if re.search(pattern, lowered):
                found.append(skill)
    # Deduplicate while preserving order
    seen = set()
    unique = []
    for s in found:
        if s not in seen:
            seen.add(s)
            unique.append(s)
    return unique


def extract_email(text: str) -> str:
    m = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", text)
    return m.group(0) if m else ""


def extract_phone(text: str) -> str:
    m = re.search(r"(\+?\d[\d\s\-\(\)]{8,}\d)", text)
    return m.group(0).strip() if m else ""


def guess_name(text: str, fallback: str) -> str:
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    for line in lines[:8]:
        # Simple heuristic: line with 2-4 capitalized words, no digits/email
        if "@" in line or any(ch.isdigit() for ch in line):
            continue
        words = line.split()
        if 2 <= len(words) <= 4 and all(w[0].isupper() for w in words if w[0].isalpha()):
            return line
    return fallback


def estimate_experience_years(text: str) -> float:
    lowered = text.lower()
    matches = re.findall(r"(\d+(?:\.\d+)?)\s*\+?\s*(?:years?|yrs?)", lowered)
    if not matches:
        return 0.0
    nums = [float(m) for m in matches if 0 < float(m) < 60]
    return max(nums) if nums else 0.0


def rank_resumes(job_description: str, resumes: List[Dict]) -> List[Dict]:
    if not resumes:
        return []

    jd_skills = find_skills(job_description, SKILL_DICTIONARY)

    corpus = [normalize_text(job_description)] + [normalize_text(r["text"]) for r in resumes]
    vectorizer = TfidfVectorizer(
        stop_words="english",
        ngram_range=(1, 2),
        max_df=0.95,
        min_df=1,
        sublinear_tf=True,
    )
    matrix = vectorizer.fit_transform(corpus)
    sims = cosine_similarity(matrix[0:1], matrix[1:]).flatten()

    results = []
    for resume, sim in zip(resumes, sims):
        resume_skills = find_skills(resume["text"], SKILL_DICTIONARY)
        matched = [s for s in jd_skills if s in resume_skills]
        missing = [s for s in jd_skills if s not in resume_skills]

        skill_coverage = (len(matched) / len(jd_skills)) if jd_skills else 0.0
        # Blend semantic similarity with skill coverage for a more useful score.
        composite = 0.65 * float(sim) + 0.35 * float(skill_coverage)

        results.append({
            "id": resume["id"],
            "filename": resume["filename"],
            "name": resume["name"],
            "email": resume["email"],
            "phone": resume["phone"],
            "experience_years": resume["experience_years"],
            "similarity": round(float(sim) * 100, 2),
            "skill_coverage": round(skill_coverage * 100, 2),
            "score": round(composite * 100, 2),
            "matched_skills": matched,
            "missing_skills": missing,
            "all_skills": resume_skills,
            "preview": resume["text"][:600],
        })

    results.sort(key=lambda r: r["score"], reverse=True)
    for idx, r in enumerate(results, start=1):
        r["rank"] = idx
    return results


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/rank", methods=["POST"])
def api_rank():
    job_description = (request.form.get("job_description") or "").strip()
    if not job_description:
        return jsonify({"error": "Job description is required."}), 400

    files = request.files.getlist("resumes")
    if not files or all(not f.filename for f in files):
        return jsonify({"error": "Please upload at least one resume."}), 400

    parsed: List[Dict] = []
    errors: List[str] = []
    for f in files:
        if not f or not f.filename:
            continue
        if not allowed_file(f.filename):
            errors.append(f"{f.filename}: unsupported file type (use PDF or DOCX).")
            continue
        try:
            stream = io.BytesIO(f.read())
            text = extract_text(f.filename, stream)
            if not text or len(text.strip()) < 20:
                errors.append(f"{f.filename}: could not extract readable text.")
                continue
            parsed.append({
                "id": uuid.uuid4().hex[:8],
                "filename": f.filename,
                "text": text,
                "name": guess_name(text, fallback=os.path.splitext(f.filename)[0]),
                "email": extract_email(text),
                "phone": extract_phone(text),
                "experience_years": estimate_experience_years(text),
            })
        except Exception as e:  # noqa: BLE001
            errors.append(f"{f.filename}: failed to parse ({e.__class__.__name__}).")

    if not parsed:
        return jsonify({"error": "No resumes could be processed.", "details": errors}), 400

    ranked = rank_resumes(job_description, parsed)
    jd_skills = find_skills(job_description, SKILL_DICTIONARY)

    return jsonify({
        "job_skills": jd_skills,
        "total": len(ranked),
        "candidates": ranked,
        "warnings": errors,
    })


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
